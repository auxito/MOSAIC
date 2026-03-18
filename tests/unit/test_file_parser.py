"""
FileParser 单元测试 — PDF / DOCX / TXT 解析 + parse_bytes。
"""

from __future__ import annotations

import io
import tempfile
from pathlib import Path

import pytest

from mosaic.resume.file_parser import FileParser
from mosaic.resume.schema import ResumeData


# ============================================================
# Fixtures
# ============================================================

@pytest.fixture
def parser():
    return FileParser()


@pytest.fixture
def sample_text() -> str:
    return "张三\n求职意向: 后端工程师\n\n## 教育经历\n清华大学 计算机科学 2018-2022"


@pytest.fixture
def txt_file(sample_text, tmp_path) -> Path:
    p = tmp_path / "resume.txt"
    p.write_text(sample_text, encoding="utf-8")
    return p


@pytest.fixture
def md_file(sample_text, tmp_path) -> Path:
    p = tmp_path / "resume.md"
    p.write_text(sample_text, encoding="utf-8")
    return p


@pytest.fixture
def gb18030_file(tmp_path) -> Path:
    """GB18030 编码的中文文件"""
    text = "李四\n求职意向: 前端工程师\n技能: JavaScript, React"
    p = tmp_path / "resume_cn.txt"
    p.write_bytes(text.encode("gb18030"))
    return p


# ============================================================
# TXT / MD 解析
# ============================================================

class TestTextParsing:
    def test_parse_txt(self, parser, txt_file, sample_text):
        result = parser.parse(txt_file)
        assert isinstance(result, ResumeData)
        assert result.name == "候选人"
        assert result.summary == sample_text.strip()

    def test_parse_md(self, parser, md_file, sample_text):
        result = parser.parse(md_file)
        assert isinstance(result, ResumeData)
        assert result.summary == sample_text.strip()

    def test_parse_gb18030_fallback(self, parser, gb18030_file):
        """测试 gb18030 编码 fallback"""
        result = parser.parse(gb18030_file)
        assert isinstance(result, ResumeData)
        assert "李四" in result.summary
        assert "JavaScript" in result.summary

    def test_parse_empty_text(self, parser, tmp_path):
        p = tmp_path / "empty.txt"
        p.write_text("", encoding="utf-8")
        result = parser.parse(p)
        assert isinstance(result, ResumeData)
        assert result.summary == ""


# ============================================================
# PDF 解析
# ============================================================

class TestPdfParsing:
    @pytest.fixture
    def pdf_file(self, tmp_path) -> Path:
        """用 pdfplumber 的底层 pdfminer 创建简单 PDF"""
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            pytest.skip("reportlab not installed, skipping PDF creation test")

        p = tmp_path / "resume.pdf"
        c = canvas.Canvas(str(p))
        c.drawString(72, 750, "Zhang San - Software Engineer")
        c.drawString(72, 730, "Skills: Python, Docker, Kubernetes")
        c.showPage()
        c.save()
        return p

    def test_parse_pdf(self, parser, pdf_file):
        result = parser.parse(pdf_file)
        assert isinstance(result, ResumeData)
        assert "Zhang San" in result.summary or "Software Engineer" in result.summary

    def test_parse_pdf_bytes(self, pdf_file):
        data = pdf_file.read_bytes()
        result = FileParser.parse_bytes("resume.pdf", data)
        assert isinstance(result, ResumeData)
        assert len(result.summary) > 0


# ============================================================
# DOCX 解析
# ============================================================

class TestDocxParsing:
    @pytest.fixture
    def docx_file(self, tmp_path) -> Path:
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed, skipping DOCX test")

        p = tmp_path / "resume.docx"
        doc = Document()
        doc.add_heading("王五 - 全栈工程师", level=1)
        doc.add_heading("教育经历", level=2)
        doc.add_paragraph("北京大学 计算机科学 2016-2020")
        doc.add_heading("工作经历", level=2)
        doc.add_paragraph("字节跳动 后端工程师 2020-2024")
        doc.save(str(p))
        return p

    def test_parse_docx(self, parser, docx_file):
        result = parser.parse(docx_file)
        assert isinstance(result, ResumeData)
        assert "王五" in result.summary
        assert "教育经历" in result.summary

    def test_parse_docx_preserves_headings(self, parser, docx_file):
        result = parser.parse(docx_file)
        # 标题应该被转换为 markdown 格式
        assert "# " in result.summary or "##" in result.summary

    def test_parse_docx_bytes(self, docx_file):
        data = docx_file.read_bytes()
        result = FileParser.parse_bytes("resume.docx", data)
        assert isinstance(result, ResumeData)
        assert "王五" in result.summary


# ============================================================
# parse_bytes
# ============================================================

class TestParseBytes:
    def test_parse_bytes_txt(self, sample_text):
        data = sample_text.encode("utf-8")
        result = FileParser.parse_bytes("resume.txt", data)
        assert isinstance(result, ResumeData)
        assert result.summary == sample_text.strip()

    def test_parse_bytes_md(self, sample_text):
        data = sample_text.encode("utf-8")
        result = FileParser.parse_bytes("resume.md", data)
        assert isinstance(result, ResumeData)
        assert "张三" in result.summary

    def test_parse_bytes_gb18030(self):
        text = "赵六 前端工程师"
        data = text.encode("gb18030")
        result = FileParser.parse_bytes("resume.txt", data)
        assert "赵六" in result.summary


# ============================================================
# 错误处理
# ============================================================

class TestErrors:
    def test_unsupported_format(self, parser, tmp_path):
        p = tmp_path / "resume.jpg"
        p.write_bytes(b"\xff\xd8\xff")
        with pytest.raises(ValueError, match="Unsupported format"):
            parser.parse(p)

    def test_unsupported_format_bytes(self):
        with pytest.raises(ValueError, match="Unsupported format"):
            FileParser.parse_bytes("resume.jpg", b"data")

    def test_file_not_found(self, parser):
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/path/resume.txt")

    def test_parse_bytes_unsupported_ext(self):
        with pytest.raises(ValueError, match="Unsupported format"):
            FileParser.parse_bytes("data.csv", b"col1,col2")
